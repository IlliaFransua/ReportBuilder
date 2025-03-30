from docx import Document
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Pt, Cm
from functools import wraps
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from TextCleaner import TextCleaner
import os
import re
from config import *


def parse_nested_text(text):
    text = text.strip()
    lines = text.splitlines()
    stack = [(0, [])]

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        indent = len(line) - len(line.lstrip())
        while stack and indent < stack[-1][0]:
            stack.pop()

        if indent == stack[-1][0]:
            stack[-1][1].append(stripped)
        else:
            new_list = [stripped]
            stack[-1][1].append(new_list)
            stack.append((indent, new_list))

    return stack[0][1]


class WordReportBuilder:
    def __init__(self, filename, section_number=1):
        self.filename = filename
        self.document = Document()
        self._create_custom_styles()

        self.section_number = section_number - 1
        self.sub_section_number = 0
        self.sub_sub_section_number = 0

        self.section_listing_number = 0
        self.sub_section_listing_number = 0
        self.sub_sub_section_listing_number = 0

        self.section_image_number = 0
        self.sub_section_image_number = 0
        self.sub_sub_section_image_number = 0

        self.section_table_number = 0
        self.sub_section_table_number = 0
        self.sub_sub_section_table_number = 0

    def _create_custom_styles(self):
        styles = self.document.styles

        self.typical_text_style = styles.add_style('practice_typical_text_style', 1)
        self.typical_text_style.font.name = 'Times New Roman'
        self.typical_text_style.font.size = Pt(14)
        self.typical_text_style.paragraph_format.first_line_indent = Cm(1.25)
        self.typical_text_style.paragraph_format.line_spacing = 1.5
        self.typical_text_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.typical_text_style.paragraph_format.space_before = Pt(0)
        self.typical_text_style.paragraph_format.space_after = Pt(0)

        self.empty_line_style = styles.add_style('practice_empty_line_style', 1)
        self.empty_line_style.font.name = 'Times New Roman'
        self.empty_line_style.font.size = Pt(14)
        self.empty_line_style.paragraph_format.first_line_indent = Cm(1.25)
        self.empty_line_style.paragraph_format.line_spacing = 1.25
        self.empty_line_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.empty_line_style.paragraph_format.space_before = Pt(0)
        self.empty_line_style.paragraph_format.space_after = Pt(0)

        self.practice_section_style = styles.add_style('practice_section_style', 1)
        self.practice_section_style.font.name = 'Times New Roman'
        self.practice_section_style.font.size = Pt(14)
        self.practice_section_style.font.all_caps = True
        self.practice_section_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.practice_section_style.paragraph_format.first_line_indent = Cm(1.25)
        self.practice_section_style.paragraph_format.line_spacing = 1.5
        self.practice_section_style.paragraph_format.space_before = Pt(0)
        self.practice_section_style.paragraph_format.space_after = Pt(0)

        self.practice_section_style = styles.add_style('practice_section_style_center', 1)
        self.practice_section_style.font.name = 'Times New Roman'
        self.practice_section_style.font.size = Pt(14)
        self.practice_section_style.font.all_caps = True
        self.practice_section_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.practice_section_style.paragraph_format.line_spacing = 1.5
        self.practice_section_style.paragraph_format.space_before = Pt(0)
        self.practice_section_style.paragraph_format.space_after = Pt(0)

        self.practice_code_style = styles.add_style('practice_code_style', 1)
        self.practice_code_style.font.name = 'Courier New'
        self.practice_code_style.font.size = Pt(10)
        self.practice_code_style.paragraph_format.line_spacing = 1
        self.practice_code_style.paragraph_format.first_line_indent = Cm(0)
        self.practice_code_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self.practice_code_style.paragraph_format.space_before = Pt(0)
        self.practice_code_style.paragraph_format.space_after = Pt(0)

    def _add_empty_line(self):
        self.document.add_paragraph(style='practice_empty_line_style')

    def _remove_section_number(self):
        self.section_number = 0

    def _remove_sub_section_number(self):
        self.sub_section_number = 0

    def _remove_sub_sub_section_number(self):
        self.sub_sub_section_number = 0

    def add_page_break(self):
        self.document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def add_section(self, text):
        self.section_number += 1
        self.sub_section_number = 0
        self.sub_sub_section_number = 0

        self.section_image_number = 0
        self.sub_section_image_number = 0
        self.sub_sub_section_image_number = 0

        self.section_table_number = 0
        self.sub_section_table_number = 0
        self.sub_sub_section_table_number = 0

        self.add_page_break()

        self.document.add_paragraph(f"{self.section_number} " + text, style='practice_section_style')

    def add_sub_section(self, text):
        self.sub_section_number += 1
        self.sub_sub_section_number = 0

        self.sub_section_image_number = 0
        self.sub_sub_section_image_number = 0

        self.sub_section_table_number = 0
        self.sub_sub_section_table_number = 0

        self.document.add_paragraph(f"{self.section_number}.{self.sub_section_number} " + text,
                                    style='practice_typical_text_style')

    def add_sub_sub_section(self, text):
        self.sub_sub_section_number += 1

        self.sub_sub_section_image_number = 0

        self.sub_sub_section_table_number = 0

        self.document.add_paragraph(
            f"{self.section_number}.{self.sub_section_number}.{self.sub_sub_section_number} " + text,
            style='practice_typical_text_style')

    def add_text(self, text):
        text = text.strip()
        lines = text.splitlines()

        for line in lines:
            line = line.strip()

            if line:
                if not line.endswith('.'):
                    line += '.'

                self.document.add_paragraph(line, style='practice_typical_text_style')

    def _add_list_of_level(self, description, items, level=1):
        def is_flat_list(items):
            return all(not isinstance(item, list) for item in items)

        self.document.add_paragraph(description, style='practice_typical_text_style')

        if is_flat_list(items):
            for item in items:
                para = self.document.add_paragraph(style='practice_typical_text_style')
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                para.paragraph_format.left_indent = Cm(0)
                para.add_run(f"– {item.strip(" -–").lower()}")
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
        else:
            LEVEL_STYLES = [
                lambda i: f"{chr(1072 + i % 32)})",  # Укр буквы: а, б, в, ...
                lambda i: f"{i + 1})",  # Арабские цифры: 1, 2, 3, ...
                lambda i: "–"
            ]
            NUM_LEVELS = len(LEVEL_STYLES)

            def add_sublist(items, current_level):
                style_func = LEVEL_STYLES[(current_level - 1) % NUM_LEVELS]
                counter = 0

                for item in items:
                    if isinstance(item, list):
                        add_sublist(item, current_level + 1)
                    else:
                        if item.strip():
                            para = self.document.add_paragraph(style='practice_typical_text_style')
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            para.paragraph_format.left_indent = Cm((current_level - 1) * 1)
                            counter += 1
                            prefix = style_func(counter - 1)

                            if prefix == "–":
                                para.add_run(f"{prefix} {item.strip(" -–").lower()}")
                            else:
                                para.add_run(f"{prefix} {item.strip(" -–").lower()}")

                            para.paragraph_format.space_after = Pt(0)
                            para.paragraph_format.space_before = Pt(0)

            add_sublist(items, level)

    def add_list(self, description, items):
        if not description.endswith(':'):
            description += ':'

        def add_colon_and_punctuation(items):
            result = []
            for i in range(len(items)):
                item = items[i]

                has_sublist = isinstance(item, list)

                if isinstance(item, str):
                    if i < len(items) - 1 and isinstance(items[i + 1], list):
                        if not item.endswith(':'):
                            result.append(f"{item.strip(" -–.")}:")
                        else:
                            result.append(item)
                    elif i < len(items) - 1 and isinstance(items[i + 1], str):
                        result.append(f"{item.strip(" -–.")};")
                    elif i == len(items) - 1:
                        result.append(f"{item.strip(" -–.")}.")
                    else:
                        result.append(item)

                elif isinstance(item, list):
                    if result and not isinstance(result[-1], list):
                        if not result[-1].endswith(':'):
                            result[-1] += ":"
                    result.append(add_colon_and_punctuation(item))

            return result

        if isinstance(items, str):
            items = parse_nested_text(items)

        items_with_colons_and_punctuation = add_colon_and_punctuation(items)

        self._add_list_of_level(description, items_with_colons_and_punctuation)

    def _add_numbered_list_of_level(self, description, items, level=1):
        def is_flat_list(items):
            return all(not isinstance(item, list) for item in items)

        self.document.add_paragraph(description, style='practice_typical_text_style')

        if is_flat_list(items):
            for index, item in enumerate(items, start=1):
                para = self.document.add_paragraph(style='practice_typical_text_style')
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                para.paragraph_format.left_indent = Cm(0)
                para.add_run(f"{index}) {item}")
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
        else:
            def add_sublist(items, current_level):
                counter = 0
                for item in items:
                    if isinstance(item, list):
                        add_sublist(item, current_level + 1)
                    else:
                        counter += 1
                        para = self.document.add_paragraph(style='practice_typical_text_style')
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        para.paragraph_format.left_indent = Cm((current_level - 1) * 1)
                        para.add_run(f"{counter}. {item}")
                        para.paragraph_format.space_after = Pt(0)
                        para.paragraph_format.space_before = Pt(0)

            add_sublist(items, level)

    def add_numbered_list(self, description, items):
        if not description.endswith(':'):
            description += ':'

        self._add_numbered_list_of_level(description, items)

    def add_introduction(self, text):
        self.add_page_break()
        self.document.add_paragraph(f"ВСТУП", style='practice_section_style_center')
        self._add_empty_line()
        self.add_text(text)

    def add_source_list(self):
        self.add_page_break()
        self.document.add_paragraph(f"СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ", style='practice_section_style_center')
        self._add_empty_line()

        global source_list

        for number, text in source_list.get_sources().items():
            para = self.document.add_paragraph(style='practice_typical_text_style')
            self.empty_line_style.paragraph_format.first_line_indent = Cm(1.25)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.add_run(f"{number}. {text}")

        self.add_page_break()

    def add_image(self, image_path, description):
        if not os.path.exists(image_path):
            base_name = os.path.basename(image_path)
            image_path = os.path.join('Images', base_name)

            if not os.path.exists(image_path):
                raise FileNotFoundError(f"Image not found: {image_path} or in 'Images' directory")

        current_image_number = self.get_current_image_number()

        self.add_text(f"{description} зображено на рисунку {current_image_number}.")
        self._add_empty_line()

        paragraph = self.document.add_paragraph(style='practice_typical_text_style')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.add_run().add_picture(image_path, width=Cm(10))

        result_description = f"Рисунок {current_image_number} – {description}"
        paragraph = self.document.add_paragraph(result_description, style='practice_typical_text_style')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_listing(self, file_path):
        try:
            file_name = os.path.basename(file_path)

            with open(file_path, 'r', encoding='utf-8') as f:
                code = f.read()

            listing_number = self.get_current_listing_number()

            self.add_text(f"Вміст файлу {file_name} відображено у листингу {listing_number}.")
            self._add_empty_line()
            self.document.add_paragraph(f"Лістинг {listing_number} – Вміст {file_name}",
                                        style='practice_typical_text_style')
            self.document.add_paragraph(code, style='practice_code_style')

        except Exception as e:
            print(f"Ошибка при обработке файла {file_path}: {e}")

    def add_table(self, description, data, current_table_number):
        result_description = "Тблиця " + current_table_number + " – " + description
        self.document.add_paragraph(f"{result_description}", style='practice_typical_text_style')

        max_columns = len(data[0])

        table = self.document.add_table(rows=len(data), cols=len(data[0]))

        for i, row in enumerate(data):
            for j in range(max_columns):
                if j < len(row):
                    table.cell(i, j).text = str(row[j])
                else:
                    table.cell(i, j).text = ""

                for paragraph in table.cell(i, j).paragraphs:
                    paragraph.paragraph_format.line_spacing = 1
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)

                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    table.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'

    def save(self):
        self.document.save(self.filename)
        print(f"Document saved: {self.filename}")

    def get_current_image_number(self):
        result_number = f"{self.section_number}"

        if self.sub_section_number:
            result_number += f".{self.sub_section_number}"

            if self.sub_sub_section_number:
                result_number += f".{self.sub_sub_section_number}"
                self.sub_sub_section_image_number += 1
                result_number += f".{self.sub_sub_section_image_number}"
            else:
                self.sub_section_image_number += 1
                result_number += f".{self.sub_section_image_number}"
        else:
            self.section_image_number += 1
            result_number += f".{self.section_image_number}"

        return result_number

    def get_current_listing_number(self):
        result_number = f"{self.section_number}"

        if self.sub_section_number:
            result_number += f".{self.sub_section_number}"

            if self.sub_sub_section_number:
                result_number += f".{self.sub_sub_section_number}"
                self.sub_sub_section_listing_number += 1
                result_number += f".{self.sub_sub_section_listing_number}"
            else:
                self.sub_section_listing_number += 1
                result_number += f".{self.sub_section_listing_number}"
        else:
            self.section_listing_number += 1
            result_number += f".{self.section_listing_number}"

        return result_number

    def get_current_table_number(self):
        result_number = f"{self.section_number}"

        if self.sub_section_number:
            result_number += f".{self.sub_section_number}"

            if self.sub_sub_section_number:
                result_number += f".{self.sub_sub_section_number}"
                self.sub_sub_section_image_number += 1
                result_number += f".{self.sub_sub_section_image_number}"
            else:
                self.sub_section_image_number += 1
                result_number += f".{self.sub_section_image_number}"
        else:
            self.section_image_number += 1
            result_number += f".{self.section_image_number}"

        return result_number
