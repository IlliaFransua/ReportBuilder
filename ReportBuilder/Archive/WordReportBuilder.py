from functools import wraps

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm


def parse_nested_text(text):
    """
    Parses a multi-line text with indentation (tabs or 4 spaces) into a structure of nested lists.
    Removes unnecessary spaces and indentation from the sides.

    :param text: Multi-line text with indentation.
    :return: Nested list.
    """
    text = text.strip()
    lines = text.splitlines()
    stack = [(0, [])]

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        indent = len(line) - len(line.lstrip())
        while stack and indent < stack[-1][0]:
            stack.pop()

        if indent == stack[-1][0]:
            stack[-1][1].append(stripped)
        else:
            new_list = [stripped]
            stack[-1][1].append(new_list)
            stack.append((indent, new_list))

    return stack[0][1]


class WordReportBuilder:
    def __init__(self, filename):
        self.filename = filename
        self.document = Document()
        self.table_count = 1
        self.current_part = 1
        self._create_custom_styles()

    def _create_custom_styles(self):
        styles = self.document.styles

        self.custom_style = styles.add_style('CustomStyle', 1)  # 1 is a paragraph style
        self.custom_style.font.name = 'Times New Roman'
        self.custom_style.font.size = Pt(14)
        self.custom_style.paragraph_format.line_spacing = 1.5
        self.custom_style.paragraph_format.first_line_indent = Cm(1.25)
        self.custom_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.custom_style.paragraph_format.space_before = Pt(0)
        self.custom_style.paragraph_format.space_after = Pt(0)

        self.blank_line_style = styles.add_style('BlankLineStyle', 1)
        self.blank_line_style.font.name = 'Times New Roman'
        self.blank_line_style.font.size = Pt(14)
        self.blank_line_style.paragraph_format.line_spacing = 1.25
        self.blank_line_style.paragraph_format.first_line_indent = Cm(1.25)
        self.blank_line_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.blank_line_style.paragraph_format.space_before = Pt(0)
        self.blank_line_style.paragraph_format.space_after = Pt(0)

        self.section_title_style = styles.add_style('SectionTitleStyle', 1)
        self.section_title_style.font.name = 'Times New Roman'
        self.section_title_style.font.size = Pt(16)
        self.section_title_style.font.all_caps = True
        self.section_title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.section_title_style.paragraph_format.line_spacing = 1
        self.section_title_style.paragraph_format.space_after = Pt(21)

        self.section_title_style = styles.add_style('PracticeSectionTitleStyle', 1)
        self.section_title_style.font.name = 'Times New Roman'
        self.section_title_style.font.size = Pt(16)
        self.section_title_style.font.all_caps = True
        self.section_title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.section_title_style.paragraph_format.line_spacing = 1
        self.section_title_style.paragraph_format.space_after = Pt(21)

        self.lab_topic_style = styles.add_style('LabTopicStyle', 1)
        self.lab_topic_style.font.name = 'Times New Roman'
        self.lab_topic_style.font.size = Pt(14)
        self.lab_topic_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.lab_topic_style.paragraph_format.line_spacing = 1.5

        self.lab_goal_style = styles.add_style('LabGoalStyle', 1)
        self.lab_goal_style.font.name = 'Times New Roman'
        self.lab_goal_style.font.size = Pt(14)
        self.lab_goal_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.lab_goal_style.paragraph_format.first_line_indent = Cm(1.25)
        self.lab_goal_style.paragraph_format.line_spacing = 1.5
        self.lab_goal_style.paragraph_format.space_before = Pt(0)
        self.lab_goal_style.paragraph_format.space_after = Pt(0)

        self.lab_process_style = styles.add_style('LabProcessStyle', 1)
        self.lab_process_style.font.name = 'Times New Roman'
        self.lab_process_style.font.size = Pt(14)
        self.lab_process_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.lab_process_style.paragraph_format.first_line_indent = Cm(0)
        self.lab_process_style.paragraph_format.line_spacing = 1.5

        self.lab_conclusion_style = styles.add_style('LabConclusionStyle', 1)
        self.lab_conclusion_style.font.name = 'Times New Roman'
        self.lab_conclusion_style.font.size = Pt(14)
        self.lab_conclusion_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.lab_conclusion_style.paragraph_format.first_line_indent = Cm(1.25)
        self.lab_conclusion_style.paragraph_format.line_spacing = 1.5

        self.code_style = styles.add_style('CodeStyle', 1)
        self.code_style.font.name = 'Courier New'
        self.code_style.font.size = Pt(10)
        self.code_style.paragraph_format.line_spacing = 1
        self.code_style.paragraph_format.first_line_indent = Cm(0)
        self.code_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self.code_style.paragraph_format.space_before = Pt(0)
        self.code_style.paragraph_format.space_after = Pt(0)

    def _add_blank_line(self):
        """
        Adds a blank line to the document.
        """
        self.document.add_paragraph(style='BlankLineStyle')

    def add_blank_line_decorator(func):
        """
        Decorator that adds a blank line after the decorated function runs.
        """

        @wraps(func)
        def wrapper(self, *args, **kwargs):
            result = func(self, *args, **kwargs)
            self._add_blank_line()
            return result

        return wrapper

    def append_text(self, text):
        """
        Adds text using the specified style.
        """
        self.document.add_paragraph(text, style='CustomStyle')

    @add_blank_line_decorator
    def add_image(self, image_path, description=""):
        """
        Adds an image centered in the document.
        """
        paragraph = self.document.add_paragraph(style='CustomStyle')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.add_run().add_picture(image_path, width=Cm(10))

        if description:
            paragraph = self.document.add_paragraph(description, style='CustomStyle')
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_section_title(self, title):
        self.document.add_paragraph(title, style='SectionTitleStyle')

    def add_lab_topic(self, topic):
        self.document.add_paragraph(topic, style='LabTopicStyle')

    @add_blank_line_decorator
    def add_lab_goal(self, goal):
        self.document.add_paragraph(goal, style='LabGoalStyle')

    def add_lab_process(self, process):
        self.document.add_paragraph(process, style='LabProcessStyle')

    @add_blank_line_decorator
    def add_listing(self, description, code):
        self.document.add_paragraph(description, style='CustomStyle')
        self.document.add_paragraph(code, style='CodeStyle')

    def add_list_of_level(self, description, items, level=1):
        def is_flat_list(items):
            """
            Checks if the list is flat (does not contain nested lists).
            """
            return all(not isinstance(item, list) for item in items)

        self.document.add_paragraph(description, style='CustomStyle')

        # If the list is flat, we add everything with a long dash
        if is_flat_list(items):
            for item in items:
                para = self.document.add_paragraph(style='CustomStyle')
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                para.paragraph_format.left_indent = Cm(0)
                para.add_run(f"– {item}")
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
        else:
            LEVEL_STYLES = [
                lambda i: f"{chr(1072 + i % 32)})",  # Ukrainian letters: а, б, в, ...
                lambda i: f"{i + 1}.",  # 1, 2, 3, ...
                lambda i: "–"
            ]
            NUM_LEVELS = len(LEVEL_STYLES)

            def add_sublist(items, current_level):
                style_func = LEVEL_STYLES[(current_level - 1) % NUM_LEVELS]
                level_counter = 0

                for item in items:
                    if isinstance(item, list):
                        add_sublist(item, current_level + 1)
                    else:
                        if item.strip():
                            para = self.document.add_paragraph(style='CustomStyle')
                            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            para.paragraph_format.left_indent = Cm(
                                (current_level - 1) * 1)
                            level_counter += 1
                            prefix = style_func(level_counter - 1)

                            if prefix == "—":
                                para.add_run(f"{prefix} {item}")
                            else:
                                para.add_run(f"{prefix} {item}")

                            para.paragraph_format.space_after = Pt(0)
                            para.paragraph_format.space_before = Pt(0)

            add_sublist(items, level)

    @add_blank_line_decorator
    def add_list(self, description, items):
        """
        Adds a list with a colon before it, as specified in the conditions.

        :param description: The description to be added before the list.
        :param items: A list of items to be included in the list.
                      This can also be a multi-line string that will be parsed.
        """

        def add_colon_and_punctuation(items):
            """
            Adds a colon, semicolon, or period to list items based on their position.

            This function processes each item in the list and determines
            the appropriate punctuation to add based on the next item.

            :param items: A list of items to process.
            :return: A list of items with the appropriate punctuation added.
            """
            result = []
            for i in range(len(items)):
                item = items[i]

                has_sublist = isinstance(item, list)

                if isinstance(item, str):
                    if i < len(items) - 1 and isinstance(items[i + 1], list):
                        if not item.endswith(':'):
                            result.append(f"{item}:")
                        else:
                            result.append(item)
                    elif i < len(items) - 1 and isinstance(items[i + 1], str):
                        result.append(f"{item};")
                    elif i == len(items) - 1:
                        result.append(f"{item}.")
                    else:
                        result.append(item)

                elif isinstance(item, list):
                    if result and not isinstance(result[-1],
                                                 list):  # If the previous item is not a sublist
                        if not result[-1].endswith(':'):
                            result[-1] += ":"
                    result.append(add_colon_and_punctuation(item))  # We handle the nested list recursively

            return result

        if isinstance(items, str):  # If multiple-line text is passed
            items = parse_nested_text(items)

        items_with_colons_and_punctuation = add_colon_and_punctuation(items)

        self.add_list_of_level(description, items_with_colons_and_punctuation)

    @add_blank_line_decorator
    def add_table(self, description, data):
        """
        Adds a description and a numbered table to the document.

        :param description: Text to be added before the table.
        :param data: A list of lists (or tuples) containing the data for the table.
        """

        self.document.add_paragraph(f"{description}", style='CustomStyle')

        max_columns = len(data[0])

        table = self.document.add_table(rows=len(data), cols=len(data[0]))

        for i, row in enumerate(data):
            for j in range(max_columns):
                if j < len(row):
                    table.cell(i, j).text = str(row[j])
                else:
                    table.cell(i, j).text = ""

                for paragraph in table.cell(i, j).paragraphs:
                    paragraph.paragraph_format.line_spacing = 1
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)

                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    table.cell(i, j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'

    def add_lab_conclusion(self, conclusion):
        self.document.add_paragraph(conclusion, style='LabConclusionStyle')

    def add_page_break(self):
        self.document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def save(self):
        self.add_page_break()
        self.document.save(self.filename)
        print(f"Document saved: {self.filename}")
